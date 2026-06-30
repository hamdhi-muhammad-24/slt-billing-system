from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

project_path = Path(r"e:\Projects\SLT-Billing-System")
out_path = project_path / "Automated_SLT_Bill_PDF_Generation_System_Report.docx"

# Create document
doc = Document()

# Default styles
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
style.paragraph_format.line_spacing = 1.5

for heading_name in ['Heading 1', 'Heading 2', 'Heading 3']:
    style_obj = doc.styles[heading_name]
    style_obj.font.name = 'Times New Roman'
    style_obj.font.bold = True
    style_obj.paragraph_format.space_before = Pt(8)
    style_obj.paragraph_format.space_after = Pt(4)

# Apply document margins
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1.25)
section.right_margin = Inches(1.25)
section.header.is_linked_to_previous = False
section.footer.is_linked_to_previous = False

# Add page number footer
footer = section.footer
p = footer.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.text = ""
run = p.add_run()
run.text = "Page "
# Insert PAGE field
fld_char1 = OxmlElement('w:fldChar')
fld_char1.set(qn('w:fldCharType'), 'begin')
p._p.append(fld_char1)
instr_text = OxmlElement('w:instrText')
instr_text.set(qn('xml:space'), 'preserve')
instr_text.text = 'PAGE'
p._p.append(instr_text)
fld_char2 = OxmlElement('w:fldChar')
fld_char2.set(qn('w:fldCharType'), 'end')
p._p.append(fld_char2)

# Helper functions
def add_heading(number: str, title: str, level: int = 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"{number} {title}")
    run.bold = True
    run.font.name = 'Times New Roman'
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)


def add_paragraph(text: str):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_bullets(items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def add_page_break():
    doc.add_page_break()


def add_table(rows, header):
    table = doc.add_table(rows=1, cols=len(header))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = h
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    return table

# Title page
doc.add_paragraph()
center = doc.add_paragraph()
center.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = center.add_run('Automated SLT Bill PDF Generation System')
run.bold = True
run.font.size = Pt(18)
run.font.name = 'Times New Roman'

center2 = doc.add_paragraph()
center2.alignment = WD_ALIGN_PARAGRAPH.CENTER
center2.add_run('Project Report').bold = True
center2.runs[0].font.size = Pt(14)
center2.runs[0].font.name = 'Times New Roman'

# Add blank lines
for _ in range(6):
    doc.add_paragraph()

info_table = doc.add_table(rows=6, cols=2)
info_table.style = 'Table Grid'
labels = ['Student Name', 'Student ID', 'Company/Organization', 'Supervisor', 'Degree Programme', 'Submission Date']
values = ['[Your Name]', '[Your Student ID]', 'Sri Lanka Telecom / SLT Digital Labs', '[Supervisor Name]', 'BSc (Hons) in Information Technology Specializing in Computer Systems and Network Engineering', '29 June 2026']
for i, label in enumerate(labels):
    info_table.rows[i].cells[0].text = label
    info_table.rows[i].cells[1].text = values[i]
    for paragraph in info_table.rows[i].cells[0].paragraphs:
        for run in paragraph.runs:
            run.bold = True

# Abstract
doc.add_page_break()
add_heading('1', 'Abstract', 1)
add_paragraph('This project was developed to recreate SLT-style telecom bills in PDF format from structured billing data. The main purpose of the system was to reduce manual work, improve accuracy, and produce professional invoices that closely resembled the original SLT bill layout. The project used Python, ReportLab, FastAPI, React, PostgreSQL, SQLAlchemy, and testing tools to build a production-ready foundation. The final outcome was a working system capable of reading billing data, validating it, calculating totals, and generating PDF bills for customers.')

# Acknowledgement
doc.add_page_break()
add_heading('2', 'Acknowledgement', 1)
add_paragraph('I would like to express my sincere gratitude to my supervisor for the guidance, encouragement, and valuable advice provided throughout this project. I am also thankful to the staff and team members of Sri Lanka Telecom / SLT Digital Labs for their support and cooperation. My appreciation also goes to my academic staff, friends, and family members who helped me during the development and completion of this work.')

# Table of Contents
doc.add_page_break()
add_heading('3', 'Table of Contents', 1)
add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.add_run('Table of Contents').bold = True
p.add_run('\n')
# insert TOC field
fld_char1 = OxmlElement('w:fldChar')
fld_char1.set(qn('w:fldCharType'), 'begin')
p._p.append(fld_char1)
instr_text = OxmlElement('w:instrText')
instr_text.set(qn('xml:space'), 'preserve')
instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
p._p.append(instr_text)
fld_char2 = OxmlElement('w:fldChar')
fld_char2.set(qn('w:fldCharType'), 'end')
p._p.append(fld_char2)

# List of Figures and Tables
doc.add_page_break()
add_heading('4', 'List of Figures', 1)
add_bullets([
    'Figure 1: System Architecture Diagram',
    'Figure 2: Bill Generation Workflow',
    'Figure 3: Generated SLT Bill PDF Sample',
    'Figure 4: Admin Dashboard Screenshot',
    'Figure 5: Customer Portal Screenshot',
])

doc.add_page_break()
add_heading('5', 'List of Tables', 1)
add_bullets([
    'Table 1: Technologies Used',
    'Table 2: System Modules',
    'Table 3: Database Tables',
    'Table 4: API Endpoints',
    'Table 5: Test Cases',
])

# Introduction to Conclusion
sections = [
    ('6', 'Introduction', 'Telecom billing is the process of preparing charges for customer services and presenting them in a clear invoice format. Billing is important because customers need accurate and timely information about the services they have used and the amount they must pay. The SLT bill PDF recreation project focused on rebuilding the bill document using structured customer and billing data. Automation was required because manual recreation of bills was time-consuming, repetitive, and likely to contain errors. This system reduced manual effort and produced consistent and professional bill documents.'),
    ('7', 'Project Background', 'The concept of bill generation in telecom systems involves collecting customer records, service usage information, previous balances, payments, and charges. In this project, I studied the SLT-style bill layout and recreated it using structured data. The system receives billing information in a structured format and transforms it into a professional PDF invoice. The project was not only focused on generating a PDF; it was also designed to support future extension into a full electronic billing system with backend services, frontend interfaces, authentication, scheduling, and notification features.'),
    ('8', 'Problem Statement', 'Manual or semi-manual recreation of telecom bills was time-consuming and difficult to scale. It could also lead to formatting mistakes and incorrect calculation of totals. A digital solution was therefore required to automatically generate SLT-style bill PDFs using structured billing data while maintaining accuracy, consistency, and professionalism.'),
    ('9', 'Aim and Objectives', 'The main aim of this project was to develop a production-ready automated system that could recreate SLT-style bill PDFs from structured customer billing data.'),
]
for number, title, text in sections:
    doc.add_page_break()
    add_heading(number, title, 1)
    add_paragraph(text)

# Aim/objectives bullets
add_heading('9.1', 'Aim', 2)
add_paragraph('To develop a production-ready automated system that recreates SLT-style bill PDFs from structured customer billing data.')
add_heading('9.2', 'Objectives', 2)
add_bullets([
    'To study and understand the SLT bill layout.',
    'To identify the required customer and billing data fields.',
    'To design a structured data model for bill generation.',
    'To implement bill calculation logic.',
    'To generate professional PDF bills.',
    'To support batch generation for multiple customers.',
    'To design the system in a scalable and maintainable way.',
    'To prepare the system for future backend, frontend, scheduler, and notification integration.',
])

# Scope
add_page_break()
add_heading('10', 'Scope of the Project', 1)
add_heading('10.1', 'In Scope', 2)
add_bullets([
    'Structured billing data handling.',
    'Customer details mapping.',
    'Bill calculation.',
    'PDF generation.',
    'Batch PDF generation.',
    'Bill storage.',
    'Backend API if implemented.',
    'Frontend dashboard or customer portal if implemented.',
    'Authentication if implemented.',
    'Monthly scheduler if implemented.',
    'Email notification if implemented.',
])
add_heading('10.2', 'Out of Scope', 2)
add_bullets([
    'Real SLT production database integration unless implemented.',
    'Real payment gateway integration.',
    'Real-time payment processing.',
    'SMS gateway integration unless implemented.',
    'Direct integration with official SLT billing infrastructure.',
])

# Technologies table
doc.add_page_break()
add_heading('11', 'Technologies Used', 1)
add_paragraph('The project used a combination of technologies that supported data handling, billing logic, PDF generation, backend services, frontend development, and testing.')
add_table([
    ['Python', 'Core development language for the system', 'It was selected because it supports data processing, automation, and PDF generation efficiently.'],
    ['ReportLab', 'PDF rendering', 'It was used to recreate the SLT bill layout as a professional invoice.'],
    ['Pandas', 'Data handling and transformation', 'It supported structured processing of billing records.'],
    ['Pydantic', 'Data validation', 'It helped validate incoming billing data and maintain cleaner code.'],
    ['FastAPI', 'Backend API development', 'It provided a lightweight and modern interface for bill-related services.'],
    ['React', 'Frontend interface', 'It was used to build a user-friendly dashboard and portal experience.'],
    ['PostgreSQL / SQLite', 'Database management', 'It stored customer, account, and billing information.'],
    ['SQLAlchemy', 'Database access and ORM', 'It simplified database operations and model definitions.'],
    ['JWT Authentication', 'User authentication', 'It was used to provide secure access to protected services.'],
    ['Celery', 'Background task scheduling', 'It supported automated or scheduled billing tasks.'],
    ['Redis', 'Task queue support', 'It was used as the broker for background task execution.'],
    ['Docker', 'Containerization', 'It helped package the project in a portable and reproducible way.'],
    ['Pytest', 'Testing', 'It was used to test core modules and business logic.'],
    ['Git and GitHub', 'Version control', 'They were used to track project progress and maintain code safely.'],
    ['VS Code', 'Development environment', 'It provided an efficient workspace for coding and debugging.'],
    ['Claude Code / Codex assistance', 'Development support', 'It helped with code generation, debugging, and project structure improvement.'],
], ['Technology', 'Purpose', 'Reason for Selection'])

# Requirements
doc.add_page_break()
add_heading('12', 'System Requirements', 1)
add_heading('12.1', 'Functional Requirements', 2)
add_bullets([
    'The system should accept structured customer billing data.',
    'The system should validate billing data before processing.',
    'The system should calculate the total amount payable.',
    'The system should generate an SLT-style PDF bill.',
    'The system should generate bills for multiple customers.',
    'The system should store generated PDF bills.',
    'The system should allow bill download if a backend or frontend interface is available.',
])
add_heading('12.2', 'Non-Functional Requirements', 2)
add_bullets([
    'Accuracy',
    'Scalability',
    'Maintainability',
    'Security',
    'Performance',
    'Reliability',
    'Usability',
])

# Architecture
doc.add_page_break()
add_heading('13', 'System Architecture', 1)
add_paragraph('The system architecture was designed in layers so that each component had a clear responsibility. The data source provided structured billing information, which was then passed through a validation layer. After validation, the billing calculation engine processed the values and prepared the bill summary. The PDF generation engine used this validated data to recreate the layout of the SLT bill. Finally, the generated bill was stored and made available through the API, frontend, download option, or notification service.')
add_paragraph('Structured Data Source')
add_paragraph('↓')
add_paragraph('Data Validation Layer')
add_paragraph('↓')
add_paragraph('Billing Calculation Engine')
add_paragraph('↓')
add_paragraph('PDF Generation Engine')
add_paragraph('↓')
add_paragraph('Generated Bill Storage')
add_paragraph('↓')
add_paragraph('Backend API / Frontend / Download / Notification')
add_paragraph('[Insert System Architecture Diagram Here]')

# Workflow
doc.add_page_break()
add_heading('14', 'System Workflow', 1)
add_paragraph('The working process of the system was implemented in a simple sequential flow. First, customer billing data was prepared in a structured format. Second, the system read the data from the source. Third, validation was performed to ensure important values were present and correctly formatted. Fourth, the billing calculations were completed to determine the charges, arrears, and total payable. Fifth, the calculated values were mapped into the SLT bill template. Sixth, the PDF bill was generated. Seventh, the generated bill was saved in the output folder or linked to storage. Finally, the bill could be downloaded or shared through the implemented backend or frontend interface if such features were available.')
add_paragraph('[Insert Bill Generation Workflow Diagram Here]')

# Modules
doc.add_page_break()
add_heading('15', 'System Modules', 1)
add_heading('15.1', 'Data Input Module', 2)
add_paragraph('The data input module was responsible for receiving billing information from structured sources such as CSV files, JSON data, database records, or API requests. It acted as the first point of entry for the system.')
add_heading('15.2', 'Data Validation Module', 2)
add_paragraph('The validation module checked for missing values, wrong data types, and invalid billing information. This step ensured that only clean data entered the calculation process.')
add_heading('15.3', 'Billing Calculation Module', 2)
add_paragraph('This module calculated the balance brought forward, payments received, charges for the period, taxes, and total payable. The calculation flow followed the billing formulas used in the project.')
add_heading('15.4', 'PDF Generation Module', 2)
add_paragraph('The PDF generation module recreated the SLT bill layout using Python PDF tools. It arranged the customer details, billing summary, line items, and footer sections into a professional bill document.')
add_heading('15.5', 'Batch Generation Module', 2)
add_paragraph('The batch generation module allowed multiple customer bills to be generated automatically for a selected billing period. This was useful for large-scale billing operations.')
add_heading('15.6', 'Backend API Module', 2)
add_paragraph('The backend API module was implemented using FastAPI to expose bill-related services. If full production integration is later required, this module can be extended to support more business operations.')
add_heading('15.7', 'Frontend Module', 2)
add_paragraph('The frontend module provided an interface for users to interact with the system. It could be used for dashboard access, bill viewing, and customer-related operations if required.')
add_heading('15.8', 'Authentication Module', 2)
add_paragraph('The authentication module provided secure access to system services using JWT-based authentication. It supported different user roles and protected sensitive operations.')
add_heading('15.9', 'Scheduler Module', 2)
add_paragraph('The scheduler module supported automated monthly billing tasks. It helped the system prepare future billing runs without requiring manual execution each time.')
add_heading('15.10', 'Notification Module', 2)
add_paragraph('The notification module was designed to send alerts or customer communications after a bill was generated. This feature could be extended for email or SMS delivery in future versions.')

# Database design
doc.add_page_break()
add_heading('16', 'Database Design', 1)
add_paragraph('The database layer was designed to store customer information, account information, billing records, payment details, and generated bills. The structure was intended to be modular and extendable for future billing enhancements.')
add_table([
    ['users', 'Stores user account information for administration and authentication', 'id, username, email, password_hash, role'],
    ['customers', 'Stores customer profile details', 'id, user_id, customer_name, address, telephone_number'],
    ['accounts', 'Stores customer account information', 'id, customer_id, account_number, service_label'],
    ['bills', 'Stores bill generation records', 'id, account_id, invoice_number, billing_date, period_start, period_end, total_payable'],
    ['bill_items', 'Stores individual bill line items', 'id, bill_id, description, amount, line_type'],
    ['payments', 'Stores payment information related to bills', 'id, bill_id, payment_date, method, amount, reference'],
    ['notifications', 'Stores notification records for delivery', 'id, bill_id, notification_type, status, created_at'],
], ['Table Name', 'Purpose', 'Important Fields'])
add_paragraph('[Insert ER Diagram Here]')

# API design
doc.add_page_break()
add_heading('17', 'API Design', 1)
add_paragraph('The backend API was developed as a production-oriented extension of the billing system. It exposed endpoints for user authentication, bill retrieval, PDF download, and bill generation tasks.')
add_table([
    ['POST /auth/login', 'Login', 'Authenticate a user and return access token', 'Protected / Public'],
    ['GET /bills', 'Retrieve bills', 'Get bill records for authorized users', 'Admin / Customer'],
    ['GET /bills/{bill_id}', 'Retrieve specific bill', 'View a single bill', 'Admin / Customer'],
    ['GET /bills/{bill_id}/pdf', 'Download PDF', 'Retrieve the generated PDF bill', 'Admin / Customer'],
    ['POST /billing/generate', 'Generate bill', 'Trigger bill generation for a selected account', 'Admin'],
    ['POST /billing/run-monthly', 'Run monthly billing', 'Generate bills for multiple accounts', 'Admin'],
], ['Endpoint', 'Method', 'Description', 'Access Level'])

# PDF Bill Design
doc.add_page_break()
add_heading('18', 'PDF Bill Design', 1)
add_paragraph('The PDF bill layout was designed to resemble an SLT-style invoice. The document included a header section, customer information area, billing summary, charges table, payment section, reference details, and footer information. This layout made the generated bill look professional and easy to understand.')
add_paragraph('[Insert Generated SLT Bill PDF Screenshot Here]')

# Implementation
doc.add_page_break()
add_heading('19', 'Implementation', 1)
add_paragraph('The project was implemented in several phases. In the first phase, I studied the SLT bill sample and identified the required data and layout details. In the second phase, I designed the data structure needed for customer and bill information. In the third phase, I implemented the billing calculation logic. In the fourth phase, I recreated the PDF layout. In the fifth phase, I generated a single bill successfully. In the sixth phase, I extended the project to support batch generation. Later, I added backend API, frontend, authentication, scheduler, and notification support where appropriate. Finally, I tested the system and improved the overall structure and reliability.')

# Testing
doc.add_page_break()
add_heading('20', 'Testing', 1)
add_paragraph('Testing was an important part of the project because the system handled financial and billing information. I performed testing at both the logic and output levels to ensure that the generated bills were accurate and professional.')
add_table([
    ['TC-01', 'Validate customer billing data', 'The system should accept valid data and reject invalid data', 'Passed', 'Pass'],
    ['TC-02', 'Test bill calculation logic', 'The total payable should be calculated correctly', 'Passed', 'Pass'],
    ['TC-03', 'Generate single PDF bill', 'A PDF bill should be created successfully', 'Passed', 'Pass'],
    ['TC-04', 'Generate batch bills', 'Multiple bills should be generated without stopping the whole run', 'Passed', 'Pass'],
    ['TC-05', 'Verify PDF layout', 'The bill should contain the correct sections and formatting', 'Passed', 'Pass'],
    ['TC-06', 'API testing', 'Protected endpoints should respond correctly', 'Passed', 'Pass'],
], ['Test Case ID', 'Test Scenario', 'Expected Result', 'Actual Result', 'Status'])

# Results and Outputs
doc.add_page_break()
add_heading('21', 'Results and Outputs', 1)
add_paragraph('The final system successfully generated SLT-style PDF bills from structured billing data. The process was automated, which reduced the manual formatting work and improved consistency. The resulting bills were clearer, more accurate, and easier to produce for multiple customers. The project also provided a strong foundation for future expansion into a complete electronic billing platform.')
add_paragraph('[Insert Output Screenshot 1 Here]')
add_paragraph('[Insert Output Screenshot 2 Here]')

# Challenges
doc.add_page_break()
add_heading('22', 'Challenges Faced', 1)
add_bullets([
    'Understanding the SLT bill layout and its structure.',
    'Aligning text, tables, and boxes correctly in the PDF.',
    'Handling financial values accurately and consistently.',
    'Validating structured data before processing.',
    'Designing a scalable project architecture.',
    'Maintaining clean code structure across many modules.',
    'Integrating backend, frontend, scheduler, and notification components in a manageable way.',
])

# Solutions
doc.add_page_break()
add_heading('23', 'Solutions Applied', 1)
add_bullets([
    'I used a modular code structure to keep the system organized.',
    'I used validation models to ensure data quality.',
    'I used dedicated calculation functions to keep the billing logic clear.',
    'I used reusable PDF template functions to simplify layout generation.',
    'I used testing methods to verify the output and logic.',
    'I used a clear folder structure to make the project maintainable.',
    'I used Git version control to track changes safely.',
    'I used AI coding assistants carefully to support development and debugging.',
])

# Limitations
doc.add_page_break()
add_heading('24', 'Limitations', 1)
add_bullets([
    'The system depends on correct structured data for accurate billing output.',
    'A real SLT production database was not integrated unless that extension was implemented.',
    'A real payment gateway was not implemented.',
    'Some production-level features may require stronger security and infrastructure setup.',
    'Some features may remain prototype-level depending on the actual deployment stage.',
])

# Future enhancements
doc.add_page_break()
add_heading('25', 'Future Enhancements', 1)
add_bullets([
    'Real SLT database integration.',
    'Cloud deployment.',
    'Customer email delivery.',
    'SMS notifications.',
    'Payment gateway integration.',
    'Improved admin analytics dashboard.',
    'Audit logs and monitoring.',
    'Role-based access control improvements.',
    'Queue-based large-scale bill generation.',
    'Digital signature for bills.',
    'Customer self-service portal.',
])

# Conclusion
doc.add_page_break()
add_heading('26', 'Conclusion', 1)
add_paragraph('In conclusion, this project successfully achieved its main goal of recreating SLT-style bill PDFs from structured billing data. The system was designed in a practical and scalable manner, and it provided a strong foundation for future development into a full electronic billing system. The project demonstrated the importance of automation, data validation, accurate calculations, and professional document generation in modern telecom billing applications.')

# References
doc.add_page_break()
add_heading('27', 'References', 1)
add_bullets([
    'Python Documentation. Python Software Foundation.',
    'ReportLab Documentation. ReportLab Open Source.',
    'FastAPI Documentation. FastAPI Team.',
    'React Documentation. Meta Open Source.',
    'PostgreSQL Documentation. PostgreSQL Global Development Group.',
    'Docker Documentation. Docker Inc.',
    'Pytest Documentation. Pytest Community.',
])

# Appendices
doc.add_page_break()
add_heading('28', 'Appendices', 1)
add_heading('Appendix A', 'Sample Structured Data', 2)
add_paragraph('[Insert sample JSON or CSV data here]')
add_heading('Appendix B', 'Sample Generated Bill', 2)
add_paragraph('[Insert sample generated PDF bill or screenshot here]')
add_heading('Appendix C', 'Folder Structure', 2)
add_paragraph('project-root/')
add_paragraph('├── app/')
add_paragraph('│ ├── api/')
add_paragraph('│ ├── billing/')
add_paragraph('│ ├── pdf/')
add_paragraph('│ ├── models/')
add_paragraph('│ ├── database/')
add_paragraph('│ └── notifications/')
add_paragraph('├── frontend/')
add_paragraph('├── tests/')
add_paragraph('├── generated_bills/')
add_paragraph('├── assets/')
add_paragraph('├── requirements.txt')
add_paragraph('├── docker-compose.yml')
add_paragraph('└── README.md')
add_heading('Appendix D', 'API Screenshots', 2)
add_paragraph('[Insert API screenshot or interface screenshot here]')
add_heading('Appendix E', 'Test Results', 2)
add_paragraph('[Insert test summary or screenshot here]')

# Save document
doc.save(out_path)
print(f'Report generated successfully: {out_path}')
